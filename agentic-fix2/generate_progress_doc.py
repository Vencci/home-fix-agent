"""Generate project progress update Word document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("Project Progress Update: Home Fix Agent", level=1)

# Student info
p = doc.add_paragraph()
p.add_run("Wenqi Wang").bold = True
p.add_run("\nBuilding Business Applications of LLMs and Generative Models")

doc.add_paragraph()

# --- Project Summary ---
doc.add_heading("Project Summary", level=2)
doc.add_paragraph(
    "Home Fix Agent is an agentic AI system that helps homeowners diagnose and fix "
    "household problems. A user uploads a photo of a broken or damaged item (e.g., a "
    "burned-out light bulb, damaged roof shingles, a leaking faucet), and a multi-agent "
    "pipeline identifies the problem, determines replacement specifications, searches for "
    "matching products, ranks them, and guides the user through ordering — all through an "
    "interactive chat interface."
)

# --- Current Prototype Status ---
doc.add_heading("Current Prototype Status", level=2)
doc.add_paragraph(
    "The prototype is fully functional and deployed live at "
    "https://home-fix-agent-ww.fly.dev/. The core agentic pipeline is complete and has been "
    "tested across multiple real-world scenarios including light bulbs, roof damage, and "
    "appliance repairs. Recent work has focused on UI polish, cost estimation, and "
    "multi-round refinement."
)

# --- Architecture ---
doc.add_heading("Architecture & Technical Stack", level=2)
doc.add_paragraph(
    "The system is built in Python using FastAPI for the web backend, with a multi-agent "
    "pipeline orchestrated through a custom incremental engine. Key components:"
)

items = [
    ("Vision Analyst Agent", "Uses GPT-4o-mini with vision capabilities to analyze uploaded photos, "
     "identify the item and problem, assess repair difficulty (1–5 scale), and list required tools."),
    ("Spec Extractor Agent", "Determines the exact replacement specifications (e.g., bulb type, wattage, "
     "base size) from the analysis, with per-field confidence scores. Supports multi-round refinement "
     "by accumulating user feedback and updating specs incrementally."),
    ("Product Searcher Agent", "Searches for matching products using SerpAPI Google Shopping integration, "
     "with intelligent fallback to mock data and LLM-generated suggestions when API is unavailable."),
    ("Product Ranker Agent", "Scores and ranks search results based on spec match, price, ratings, and "
     "availability, returning the top 5 recommendations with explanations."),
    ("Cost Estimator Agent", "Estimates full repair cost by combining parts (from real product prices), "
     "tools (with per-tool price estimates and ownership likelihood), and optional labor into a total range."),
    ("Order Manager Agent", "Handles product selection and order confirmation flow."),
]
for name, desc in items:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(name + ": ").bold = True
    p.add_run(desc)

# --- Key Features ---
doc.add_heading("Key Features Implemented", level=2)
features = [
    "Modern, polished web UI — redesigned with a hero welcome page, \"How it works\" steps, "
    "trust badges, and a modern indigo/purple color system with smooth animations",
    "Chat-driven interface — the entire diagnosis-to-order flow is presented as a natural conversation",
    "Interactive clarification — when the system is uncertain about specs, it asks the user targeted "
    "questions before (or alongside) searching",
    "Multi-round refinement loop — after seeing product results, users can provide multiple rounds "
    "of feedback (e.g., \"I need dimmable\", then \"show me cheaper options\"). The system accumulates "
    "all prior feedback, shows spec diffs after each round, and progressively improves results",
    "Full repair cost estimation — a dedicated Cost Estimator agent combines parts, tools (with "
    "per-tool breakdown and ownership likelihood), and optional labor into a total cost range displayed "
    "as a visual card in the chat",
    "Enhanced product cards — horizontal layout with inline images, visual star ratings, match score "
    "progress bars, and prominent recommendation reasons",
    "DIY vs. Hire assessment — the vision agent rates repair difficulty and recommends whether to "
    "DIY or hire a professional, with labor price estimates",
    "Real product search via SerpAPI (Google Shopping), with intelligent fallback to mock data "
    "and LLM-generated suggestions when API is unavailable",
    "Session persistence — all sessions are saved with full conversation transcripts and can be "
    "replayed from the history sidebar",
    "Deployed live on Fly.io with Docker containerization",
]
for f in features:
    doc.add_paragraph(f, style="List Bullet")

# --- Demo Scenarios Tested ---
doc.add_heading("Demo Scenarios Tested", level=2)
doc.add_paragraph(
    "The system has been tested with 20+ real sessions covering diverse home repair scenarios, "
    "including:"
)
scenarios = [
    "Light bulb replacement (identifying bulb type, base, wattage from photo)",
    "Roof shingle damage (identifying material, estimating coverage area)",
    "Various household fixtures and appliances",
    "Multi-round refinement workflows (e.g., adding dimmable requirement, adjusting price range)",
]
for s in scenarios:
    doc.add_paragraph(s, style="List Bullet")

# --- What's Next ---
doc.add_heading("Remaining Work for Final Showcase", level=2)
remaining = [
    "Enhance product search with additional retailer APIs beyond SerpAPI for broader coverage",
    "Prepare a compelling live demo walkthrough for the showcase presentation",
    "Add user accounts and saved preferences for returning users",
]
for r in remaining:
    doc.add_paragraph(r, style="List Bullet")

# --- Link ---
doc.add_heading("Links", level=2)
doc.add_paragraph("GitHub Repository: https://github.com/c8872762-tech/agentic-fix2")
doc.add_paragraph("Live Demo: https://home-fix-agent-ww.fly.dev/")

out = r"C:\Users\wenqwan\Code\agentic-fix2\Project_Progress_Update.docx"
doc.save(out)
print(f"Saved to {out}")
